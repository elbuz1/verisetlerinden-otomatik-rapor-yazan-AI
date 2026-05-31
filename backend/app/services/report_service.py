import os
import uuid
from datetime import datetime
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image, PageBreak, KeepTogether
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from app.config import get_settings

settings = get_settings()

PRIMARY_COLOR = HexColor("#6366f1")
SECONDARY_COLOR = HexColor("#8b5cf6")
ACCENT_COLOR = HexColor("#10b981")
DARK_COLOR = HexColor("#1e1b4b")
LIGHT_BG = HexColor("#f8fafc")
TEXT_COLOR = HexColor("#334155")


class ReportService:

    def __init__(self):
        self.reports_dir = settings.REPORTS_DIR
        os.makedirs(self.reports_dir, exist_ok=True)

    def generate_pdf(self, analysis: dict, ai_comments: dict, charts: dict,
                     dataset_name: str, basic_info: dict) -> str:
        report_id = str(uuid.uuid4())[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"rapor_{timestamp}_{report_id}.pdf"
        filepath = os.path.join(self.reports_dir, filename)

        doc = SimpleDocTemplate(
            filepath, pagesize=A4,
            rightMargin=2 * cm, leftMargin=2 * cm,
            topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        self._register_custom_styles(styles)
        story = []

        self._add_cover_page(story, styles, dataset_name)
        story.append(PageBreak())

        self._add_toc(story, styles)
        story.append(PageBreak())

        self._add_executive_summary(story, styles, ai_comments)
        story.append(PageBreak())

        self._add_data_overview(story, styles, analysis, basic_info)
        story.append(PageBreak())

        self._add_statistical_analysis(story, styles, analysis)
        story.append(PageBreak())

        self._add_charts_section(story, styles, charts)
        story.append(PageBreak())

        self._add_trends_section(story, styles, ai_comments)
        story.append(PageBreak())

        self._add_anomaly_section(story, styles, ai_comments, analysis)
        story.append(PageBreak())

        self._add_ai_comments_section(story, styles, ai_comments)
        story.append(PageBreak())

        self._add_recommendations(story, styles, ai_comments)
        story.append(PageBreak())

        self._add_conclusion(story, styles, ai_comments, analysis)

        doc.build(story)
        return filepath

    def _register_custom_styles(self, styles):
        styles.add(ParagraphStyle(
            name="CoverTitle",
            fontSize=28,
            leading=34,
            alignment=TA_CENTER,
            textColor=DARK_COLOR,
            spaceAfter=20,
            fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="CoverSubtitle",
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            textColor=TEXT_COLOR,
            spaceAfter=10,
        ))
        styles.add(ParagraphStyle(
            name="SectionTitle",
            fontSize=18,
            leading=24,
            textColor=DARK_COLOR,
            spaceAfter=12,
            spaceBefore=20,
            fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="SubsectionTitle",
            fontSize=13,
            leading=17,
            textColor=PRIMARY_COLOR,
            spaceAfter=8,
            spaceBefore=14,
            fontName="Helvetica-Bold",
        ))
        styles.add(ParagraphStyle(
            name="BodyText2",
            fontSize=10,
            leading=15,
            alignment=TA_JUSTIFY,
            textColor=TEXT_COLOR,
            spaceAfter=8,
        ))
        styles.add(ParagraphStyle(
            name="InsightBox",
            fontSize=10,
            leading=15,
            alignment=TA_LEFT,
            textColor=HexColor("#1e40af"),
            backColor=HexColor("#eff6ff"),
            borderPadding=10,
            spaceAfter=10,
        ))

    def _add_cover_page(self, story, styles, dataset_name: str):
        story.append(Spacer(1, 4 * cm))

        line_data = [["" for _ in range(1)]]
        line_table = Table(line_data, colWidths=[16 * cm])
        line_table.setStyle(TableStyle([
            ("LINEABOVE", (0, 0), (-1, 0), 3, PRIMARY_COLOR),
        ]))
        story.append(line_table)
        story.append(Spacer(1, 1 * cm))

        story.append(Paragraph("AI Destekli Veri Analiz Raporu", styles["CoverTitle"]))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(f"Veri Seti: {dataset_name}", styles["CoverSubtitle"]))
        story.append(Spacer(1, 0.5 * cm))

        now = datetime.now()
        story.append(Paragraph(
            f"Rapor Tarihi: {now.strftime('%d.%m.%Y %H:%M')}",
            styles["CoverSubtitle"]
        ))
        story.append(Spacer(1, 3 * cm))

        line_table2 = Table(line_data, colWidths=[16 * cm])
        line_table2.setStyle(TableStyle([
            ("LINEABOVE", (0, 0), (-1, 0), 1, SECONDARY_COLOR),
        ]))
        story.append(line_table2)
        story.append(Spacer(1, 0.5 * cm))

        story.append(Paragraph(
            "Bu rapor, AI Destekli Otomatik Rapor Yazma Sistemi tarafindan olusturulmustur. "
            "Icerisindeki analizler, istatistiksel yontemler ve kural tabanli yapay zeka "
            "mantigi kullanilarak uretilmistir.",
            ParagraphStyle("FooterNote", fontSize=8, leading=11,
                           alignment=TA_CENTER, textColor=HexColor("#94a3b8"))
        ))

    def _add_toc(self, story, styles):
        story.append(Paragraph("Icindekiler", styles["SectionTitle"]))
        story.append(Spacer(1, 0.5 * cm))

        sections = [
            "1. Yonetici Ozeti",
            "2. Veri Seti Genel Bakis",
            "3. Istatistiksel Analiz",
            "4. Grafikler ve Gorsellestirme",
            "5. Trend Analizi",
            "6. Anomali Tespiti",
            "7. AI Yorumlari ve Icgoruler",
            "8. Oneriler",
            "9. Sonuc",
        ]
        for section in sections:
            story.append(Paragraph(f"  {section}", styles["BodyText2"]))

    def _add_executive_summary(self, story, styles, ai_comments: dict):
        story.append(Paragraph("1. Yonetici Ozeti", styles["SectionTitle"]))
        story.append(Spacer(1, 0.3 * cm))

        summary = ai_comments.get("executive_summary", "")
        for paragraph in summary.split("\n\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles["BodyText2"]))
                story.append(Spacer(1, 0.2 * cm))

    def _add_data_overview(self, story, styles, analysis: dict, basic_info: dict):
        story.append(Paragraph("2. Veri Seti Genel Bakis", styles["SectionTitle"]))
        basic = analysis.get("basic_stats", {})

        overview_data = [
            ["Metrik", "Deger"],
            ["Toplam Satir", f"{basic.get('total_rows', 0):,}"],
            ["Toplam Sutun", str(basic.get("total_columns", 0))],
            ["Sayisal Sutun", str(basic.get("numeric_columns", 0))],
            ["Kategorik Sutun", str(basic.get("categorical_columns", 0))],
            ["Tekrarlayan Satir", str(basic.get("duplicate_rows", 0))],
            ["Bellek Kullanimi", f"{basic.get('memory_usage_mb', 0)} MB"],
            ["Dosya Tipi", basic_info.get("file_type", "N/A")],
            ["Dosya Boyutu", basic_info.get("file_size_readable", "N/A")],
        ]

        table = Table(overview_data, colWidths=[8 * cm, 8 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), PRIMARY_COLOR),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), LIGHT_BG]),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(table)

        quality = analysis.get("data_quality_score", {})
        if quality:
            story.append(Spacer(1, 0.5 * cm))
            story.append(Paragraph("Veri Kalitesi", styles["SubsectionTitle"]))
            quality_data = [
                ["Kriter", "Skor"],
                ["Tamlik", f"{quality.get('completeness', 0):.1f}/100"],
                ["Teklik", f"{quality.get('uniqueness', 0):.1f}/100"],
                ["Tutarlilik", f"{quality.get('consistency', 0):.1f}/100"],
                ["Genel Skor", f"{quality.get('overall', 0):.1f}/100 ({quality.get('label', '')})"],
            ]
            qt = Table(quality_data, colWidths=[8 * cm, 8 * cm])
            qt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), ACCENT_COLOR),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), LIGHT_BG]),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]))
            story.append(qt)

    def _add_statistical_analysis(self, story, styles, analysis: dict):
        story.append(Paragraph("3. Istatistiksel Analiz", styles["SectionTitle"]))

        basic = analysis.get("basic_stats", {})
        columns_stats = basic.get("columns", {})

        if columns_stats:
            story.append(Paragraph("Tanimlayici Istatistikler", styles["SubsectionTitle"]))

            for col_name, col_stats in list(columns_stats.items())[:10]:
                story.append(Paragraph(f"{col_name}", styles["SubsectionTitle"]))
                stat_data = [
                    ["Istatistik", "Deger"],
                    ["Ortalama", f"{col_stats.get('mean', 0):,.4f}"],
                    ["Medyan", f"{col_stats.get('median', 0):,.4f}"],
                    ["Std. Sapma", f"{col_stats.get('std', 0):,.4f}"],
                    ["Minimum", f"{col_stats.get('min', 0):,.4f}"],
                    ["Maksimum", f"{col_stats.get('max', 0):,.4f}"],
                    ["Q1 (25%)", f"{col_stats.get('q25', 0):,.4f}"],
                    ["Q3 (75%)", f"{col_stats.get('q75', 0):,.4f}"],
                    ["Carpiklik", f"{col_stats.get('skewness', 0):.4f}"],
                    ["Basiklik", f"{col_stats.get('kurtosis', 0):.4f}"],
                    ["Degisim Kats. (%)", f"{col_stats.get('cv', 0):.2f}"],
                ]
                st = Table(stat_data, colWidths=[6 * cm, 6 * cm])
                st.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), SECONDARY_COLOR),
                    ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#ffffff")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#e2e8f0")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#ffffff"), LIGHT_BG]),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]))
                story.append(st)
                story.append(Spacer(1, 0.3 * cm))

    def _add_charts_section(self, story, styles, charts: dict):
        story.append(Paragraph("4. Grafikler ve Gorsellestirme", styles["SectionTitle"]))

        chart_titles = {
            "overview": "Genel Bakis",
            "distribution": "Dagilim Grafikleri",
            "correlation": "Korelasyon Matrisi",
            "boxplot": "Kutu Grafikleri",
            "trend": "Trend Grafikleri",
            "category": "Kategori Grafikleri",
            "missing": "Eksik Veri Grafigi",
            "quality": "Kalite Skoru",
        }

        for key, title in chart_titles.items():
            chart_path = charts.get(key)
            if not chart_path:
                continue

            paths = chart_path if isinstance(chart_path, list) else [chart_path]
            for p in paths:
                if p and os.path.exists(p):
                    story.append(Paragraph(title, styles["SubsectionTitle"]))
                    try:
                        img = Image(p, width=15 * cm, height=10 * cm)
                        img.hAlign = "CENTER"
                        story.append(img)
                    except Exception:
                        pass
                    story.append(Spacer(1, 0.5 * cm))

    def _add_trends_section(self, story, styles, ai_comments: dict):
        story.append(Paragraph("5. Trend Analizi", styles["SectionTitle"]))

        trend_comments = ai_comments.get("trend_comments", [])
        if not trend_comments:
            story.append(Paragraph(
                "Veri setinde anlamli trend tespit edilememistir.",
                styles["BodyText2"]
            ))
            return

        for tc in trend_comments:
            story.append(Paragraph(tc["comment"], styles["BodyText2"]))
            story.append(Spacer(1, 0.2 * cm))

    def _add_anomaly_section(self, story, styles, ai_comments: dict, analysis: dict):
        story.append(Paragraph("6. Anomali Tespiti", styles["SectionTitle"]))

        anomaly_comments = ai_comments.get("anomaly_comments", [])
        if not anomaly_comments:
            story.append(Paragraph(
                "Veri setinde kritik seviyede anomali tespit edilememistir.",
                styles["BodyText2"]
            ))
            return

        for ac in anomaly_comments:
            severity_label = {"dusuk": "Dusuk", "orta": "Orta", "yuksek": "Yuksek"}.get(
                ac.get("severity", ""), ""
            )
            story.append(Paragraph(
                f"[{severity_label} Oncelik] {ac['comment']}",
                styles["BodyText2"]
            ))
            story.append(Spacer(1, 0.2 * cm))

    def _add_ai_comments_section(self, story, styles, ai_comments: dict):
        story.append(Paragraph("7. AI Yorumlari ve Icgoruler", styles["SectionTitle"]))

        story.append(Paragraph("Korelasyon Bulgulari", styles["SubsectionTitle"]))
        corr_comments = ai_comments.get("correlation_comments", [])
        if corr_comments:
            for cc in corr_comments:
                story.append(Paragraph(cc["comment"], styles["BodyText2"]))
        else:
            story.append(Paragraph("Guclu korelasyon tespit edilememistir.", styles["BodyText2"]))

        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("Dagilim Bulgulari", styles["SubsectionTitle"]))
        dist_comments = ai_comments.get("distribution_comments", [])
        if dist_comments:
            for dc in dist_comments:
                story.append(Paragraph(dc["comment"], styles["BodyText2"]))
        else:
            story.append(Paragraph("Onemli dagilim bulgulari mevcut degildir.", styles["BodyText2"]))

        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph("Eksik Veri Degerlendirmesi", styles["SubsectionTitle"]))
        missing_comment = ai_comments.get("missing_data_comment", "")
        if missing_comment:
            story.append(Paragraph(missing_comment, styles["BodyText2"]))

    def _add_recommendations(self, story, styles, ai_comments: dict):
        story.append(Paragraph("8. Oneriler", styles["SectionTitle"]))

        recommendations = ai_comments.get("recommendations", [])
        if not recommendations:
            story.append(Paragraph("Su an icin ozel bir oneri bulunmamaktadir.", styles["BodyText2"]))
            return

        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(
                f"{i}. [{rec.get('category', '')}] {rec.get('recommendation', '')}",
                styles["BodyText2"]
            ))
            story.append(Spacer(1, 0.2 * cm))

    def _add_conclusion(self, story, styles, ai_comments: dict, analysis: dict):
        story.append(Paragraph("9. Sonuc", styles["SectionTitle"]))

        basic = analysis.get("basic_stats", {})
        quality = analysis.get("data_quality_score", {})

        conclusion = (
            f"Bu raporda {basic.get('total_rows', 0):,} satirlik veri seti uzerinde "
            f"kapsamli bir analiz gerceklestirilmistir. "
            f"Veri kalitesi {quality.get('overall', 0):.0f}/100 ({quality.get('label', '')}) "
            f"olarak degerlendirilmistir.\n\n"
            f"Analiz surecinde temel istatistikler, korelasyon analizi, trend tespiti, "
            f"anomali algilama ve dagilim analizi gerceklestirilmis; "
            f"sonuclar kural tabanli yapay zeka sistemi tarafindan yorumlanmistir.\n\n"
            f"Yukaridaki oneriler dikkate alinarak, veri kalitesinin arttirilmasi ve "
            f"stratejik kararlarin desteklenmesi mumkun olacaktir."
        )

        for p in conclusion.split("\n\n"):
            story.append(Paragraph(p.strip(), styles["BodyText2"]))
            story.append(Spacer(1, 0.2 * cm))

        story.append(Spacer(1, 1 * cm))
        story.append(Paragraph(
            f"Rapor Olusturma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}",
            ParagraphStyle("Timestamp", fontSize=8, alignment=TA_CENTER, textColor=HexColor("#94a3b8"))
        ))

    def generate_docx(self, analysis: dict, ai_comments: dict, charts: dict,
                      dataset_name: str, basic_info: dict) -> str:
        report_id = str(uuid.uuid4())[:8]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"rapor_{timestamp}_{report_id}.docx"
        filepath = os.path.join(self.reports_dir, filename)

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        for i in range(1, 4):
            heading_style = doc.styles[f"Heading {i}"]
            heading_style.font.color.rgb = RGBColor(0x1e, 0x1b, 0x4b)

        doc.add_paragraph("")
        doc.add_paragraph("")
        title = doc.add_heading("AI Destekli Veri Analiz Raporu", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f"Veri Seti: {dataset_name}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p = doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        doc.add_heading("Icindekiler", level=1)
        sections = [
            "1. Yonetici Ozeti", "2. Veri Seti Genel Bakis", "3. Istatistiksel Analiz",
            "4. Grafikler", "5. Trend Analizi", "6. Anomali Tespiti",
            "7. AI Yorumlari", "8. Oneriler", "9. Sonuc",
        ]
        for s in sections:
            doc.add_paragraph(s, style="List Number")
        doc.add_page_break()

        doc.add_heading("1. Yonetici Ozeti", level=1)
        summary = ai_comments.get("executive_summary", "")
        for p in summary.split("\n\n"):
            if p.strip():
                doc.add_paragraph(p.strip())

        doc.add_page_break()
        doc.add_heading("2. Veri Seti Genel Bakis", level=1)
        basic = analysis.get("basic_stats", {})
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metrik"
        hdr_cells[1].text = "Deger"
        rows_data = [
            ("Toplam Satir", f"{basic.get('total_rows', 0):,}"),
            ("Toplam Sutun", str(basic.get("total_columns", 0))),
            ("Sayisal Sutun", str(basic.get("numeric_columns", 0))),
            ("Kategorik Sutun", str(basic.get("categorical_columns", 0))),
            ("Tekrarlayan Satir", str(basic.get("duplicate_rows", 0))),
        ]
        for metric, value in rows_data:
            row = table.add_row()
            row.cells[0].text = metric
            row.cells[1].text = value

        doc.add_page_break()
        doc.add_heading("3. Istatistiksel Analiz", level=1)
        columns_stats = basic.get("columns", {})
        for col_name, stats_data in list(columns_stats.items())[:8]:
            doc.add_heading(col_name, level=2)
            stat_table = doc.add_table(rows=1, cols=2)
            stat_table.style = "Light List Accent 1"
            hdr = stat_table.rows[0].cells
            hdr[0].text = "Istatistik"
            hdr[1].text = "Deger"
            for key, label in [
                ("mean", "Ortalama"), ("median", "Medyan"), ("std", "Std. Sapma"),
                ("min", "Minimum"), ("max", "Maksimum"),
                ("skewness", "Carpiklik"), ("kurtosis", "Basiklik"),
            ]:
                row = stat_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = f"{stats_data.get(key, 0):,.4f}"

        doc.add_page_break()
        doc.add_heading("4. Grafikler", level=1)
        chart_titles = {
            "overview": "Genel Bakis", "distribution": "Dagilim",
            "correlation": "Korelasyon", "boxplot": "Kutu Grafikleri",
            "trend": "Trend", "category": "Kategoriler",
            "missing": "Eksik Veri", "quality": "Kalite",
        }
        for key, title in chart_titles.items():
            chart_path = charts.get(key)
            if not chart_path:
                continue
            paths = chart_path if isinstance(chart_path, list) else [chart_path]
            for p in paths:
                if p and os.path.exists(p):
                    doc.add_heading(title, level=2)
                    try:
                        doc.add_picture(p, width=Inches(6))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass

        doc.add_page_break()
        doc.add_heading("5. Trend Analizi", level=1)
        for tc in ai_comments.get("trend_comments", []):
            doc.add_paragraph(tc["comment"])

        doc.add_page_break()
        doc.add_heading("6. Anomali Tespiti", level=1)
        for ac in ai_comments.get("anomaly_comments", []):
            doc.add_paragraph(f"[{ac.get('severity', '').upper()}] {ac['comment']}")

        doc.add_page_break()
        doc.add_heading("7. AI Yorumlari", level=1)
        doc.add_heading("Korelasyon Bulgulari", level=2)
        for cc in ai_comments.get("correlation_comments", []):
            doc.add_paragraph(cc["comment"])
        doc.add_heading("Dagilim Bulgulari", level=2)
        for dc in ai_comments.get("distribution_comments", []):
            doc.add_paragraph(dc["comment"])
        doc.add_heading("Eksik Veri", level=2)
        doc.add_paragraph(ai_comments.get("missing_data_comment", ""))

        doc.add_page_break()
        doc.add_heading("8. Oneriler", level=1)
        for i, rec in enumerate(ai_comments.get("recommendations", []), 1):
            doc.add_paragraph(
                f"{i}. [{rec.get('category', '')}] {rec.get('recommendation', '')}",
                style="List Number"
            )

        doc.add_page_break()
        doc.add_heading("9. Sonuc", level=1)
        quality = analysis.get("data_quality_score", {})
        doc.add_paragraph(
            f"Bu raporda {basic.get('total_rows', 0):,} satirlik veri seti uzerinde "
            f"kapsamli bir analiz gerceklestirilmistir. "
            f"Veri kalitesi {quality.get('overall', 0):.0f}/100 olarak degerlendirilmistir."
        )
        doc.add_paragraph(
            "Analiz surecinde temel istatistikler, korelasyon analizi, trend tespiti, "
            "anomali algilama ve dagilim analizi yapilmistir."
        )

        doc.save(filepath)
        return filepath


report_service = ReportService()
